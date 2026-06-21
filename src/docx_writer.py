from collections import Counter
from pathlib import Path
from typing import Optional
import re

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Emu
from docx.text.paragraph import Paragraph

from PIL import Image as _PILImage

from config import (
    REQUIRED_HEADING_FONT,
    REQUIRED_HEADING_SIZE_PT,
    REQUIRED_HEADING_BOLD,
    REQUIRED_BODY_FONT,
    REQUIRED_BODY_SIZE_PT,
    COMMENT_AUTHOR,
    COMMENT_INITIALS,
    DIAGRAM_IMAGE_WIDTH_INCHES,
    ADVANCED_VISUAL_WIDTH_INCHES,
)
from models import ReviewIssue
from utils import clean_spacing, normalize_text

# Absolute ceiling for any inserted image — safe for A4 and Letter with standard margins
MAX_IMAGE_WIDTH_INCHES = 6.0
# Absolute ceiling on HEIGHT for any inserted image. Without this, a
# naturally tall/narrow diagram (e.g. a vertical flowchart) gets its
# width forced to a fixed value and its height auto-scales proportionally
# to preserve aspect ratio — which can stretch it to 20-30+ inches tall,
# spanning many pages on its own. A typical content page has roughly
# 9-9.5in of usable vertical space (Letter/A4 minus margins, header,
# footer); capping height to a safe fraction of that keeps any single
# diagram from dominating multiple pages.
MAX_IMAGE_HEIGHT_INCHES = 4.5

# Heading font-size standard (pt), keyed by heading level.
# Heading 1 -> 12pt, Heading 2 -> 10pt, Heading 3 (and deeper) -> 10pt.
HEADING_SIZE_BY_LEVEL = {1: 12, 2: 10, 3: 10}
DEFAULT_HEADING_SIZE = 10  # used for any heading level beyond 3

# Font size (pt) applied to bold "pseudo-headings" — paragraphs that are
# NOT a real Word Heading style but function as one visually (bold,
# short, standalone line, e.g. a title-page line or an in-body bold
# label). Applies document-wide, independent of HEADING_SIZE_BY_LEVEL,
# which only governs real Heading 1/2/3 styled paragraphs.
PSEUDO_HEADING_SIZE_PT = 14
PSEUDO_HEADING_MAX_WORDS = 20

# Alignments that must NEVER be overwritten by automated formatting passes.
# A paragraph that already has one of these is considered "intentionally
# aligned" and is left exactly as-is.
PRESERVED_ALIGNMENTS = (
    WD_ALIGN_PARAGRAPH.CENTER,
    WD_ALIGN_PARAGRAPH.RIGHT,
    WD_ALIGN_PARAGRAPH.DISTRIBUTE,
)

_INLINE_MCQ_PATTERN = re.compile(
    r'[A-D][)\.].*[A-D][)\.].*[A-D][\.\)]',
    re.IGNORECASE
)
# ─────────────────────────────────────────────────────────────────────────────
# Font helpers
# ─────────────────────────────────────────────────────────────────────────────

def _get_dominant_font(doc: Document, is_heading: bool):
    """
    Find the most commonly used font name and size for heading or body paragraphs.
    Skips TOC entries — they use a different style by design and should not
    influence the dominant font calculation.
    Falls back to Normal style definition if no run-level fonts are found.
    """
    font_names = []
    font_sizes = []

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else "Normal"
        style_lower = style_name.lower()

        # Skip TOC entries — never include in dominant font calculation
        if style_lower.startswith("toc"):
            continue

        para_is_heading = style_lower.startswith("heading")
        if para_is_heading != is_heading:
            continue
        if not para.text.strip():
            continue

        for run in para.runs:
            if run.text.strip():
                if run.font.name:
                    font_names.append(run.font.name)
                if run.font.size:
                    font_sizes.append(run.font.size.pt)

    # Fallback to style definition if no explicit run-level fonts found
    if not font_names and not is_heading:
        try:
            normal_style = doc.styles["Normal"]
            if normal_style.font and normal_style.font.name:
                font_names = [normal_style.font.name]
            if normal_style.font and normal_style.font.size:
                font_sizes = [normal_style.font.size.pt]
        except Exception:
            pass

    dominant_font = Counter(font_names).most_common(1)[0][0] if font_names else None
    dominant_size = Counter(font_sizes).most_common(1)[0][0] if font_sizes else None
    return dominant_font, dominant_size


# ─────────────────────────────────────────────────────────────────────────────
# Image sizing helpers
# ─────────────────────────────────────────────────────────────────────────────

def _get_page_body_width_inches(doc: Document) -> float:
    """
    Return the usable body width in inches from the first section's page settings.
    Falls back to 6.0 inches if section info is unavailable.
    """
    try:
        section = doc.sections[0]
        page_width   = section.page_width    # Emu
        left_margin  = section.left_margin   # Emu
        right_margin = section.right_margin  # Emu
        body_width_emu = page_width - left_margin - right_margin
        return body_width_emu / 914400       # Emu → inches
    except Exception:
        return 6.0


def _safe_image_width(doc: Document, requested_inches: float) -> Inches:
    """
    Return an Inches value that fits within the page body width.
    Never exceeds the usable page area regardless of what the caller requests.

    NOTE: this only caps width. Setting width alone on a naturally tall/
    narrow image (low width:height aspect ratio) causes python-docx to
    scale height up proportionally with no limit, which can produce a
    single image 20-30+ inches tall. Use _safe_image_dimensions() instead
    wherever the image's native aspect ratio is unknown or could be
    extreme — it caps both width AND height, never width alone.
    """
    body_width = _get_page_body_width_inches(doc)
    capped = min(requested_inches, body_width, MAX_IMAGE_WIDTH_INCHES)
    return Inches(capped)


# DPI that Graphviz renders generated diagrams at (see advanced_visual_renderer.py
# and diagram_generator.py, both call dot.attr(..., dpi="180")). Used to convert
# a diagram's native pixel size back to its TRUE physical size as the renderer
# intended, rather than always stretching every diagram to one fixed width.
DIAGRAM_RENDER_DPI = 180


def _safe_image_dimensions(doc: Document, image_path: str, requested_width_inches: float):
    """
    Return (width, height) as Inches objects, sized PROPORTIONALLY to the
    image's actual content rather than always stretched to a fixed target
    width.

    Why: forcing every diagram to the same target width (e.g. 5.5in)
    regardless of its real complexity means a simple 3-node diagram and a
    dense 9-node diagram both render at an identical size — and worse, a
    naturally tall/narrow diagram (e.g. a 7-step vertical flowchart) gets
    its height inflated to match, producing a multi-page-tall image for
    content that was only ever meant to be small.

    Instead:
    1. Start from the diagram's TRUE native size — its pixel dimensions
       divided by DIAGRAM_RENDER_DPI, the resolution Graphviz actually
       rendered it at. This is the size the renderer intended.
    2. Only ever SCALE DOWN from that native size, never up — if the
       diagram is already smaller than the page bounds, it stays exactly
       that size rather than being stretched to fill a target width.
    3. If the native size exceeds either the page body width or
       MAX_IMAGE_HEIGHT_INCHES, scale down uniformly (preserving aspect
       ratio) by whichever dimension is more constraining, so oversized
       diagrams still shrink to fit the page.

    Falls back to width-only sizing (scaled down to requested_width_inches,
    no height cap) if the image's native dimensions can't be read.
    """
    max_width = min(_get_page_body_width_inches(doc), MAX_IMAGE_WIDTH_INCHES, requested_width_inches)

    try:
        with _PILImage.open(image_path) as img:
            native_w_px, native_h_px = img.size
        if native_w_px <= 0 or native_h_px <= 0:
            raise ValueError("non-positive native image dimensions")
    except Exception:
        # Can't read native dimensions — fall back to the previous
        # width-only behavior rather than failing the whole insertion.
        return Inches(max_width), None

    native_width_in = native_w_px / DIAGRAM_RENDER_DPI
    native_height_in = native_h_px / DIAGRAM_RENDER_DPI

    # Scale factor needed to bring the native size within EACH bound.
    # A value >= 1 means that bound isn't a constraint (no upscaling).
    width_scale = min(1.0, max_width / native_width_in) if native_width_in > 0 else 1.0
    height_scale = min(1.0, MAX_IMAGE_HEIGHT_INCHES / native_height_in) if native_height_in > 0 else 1.0

    # Use the smaller (more constraining) scale factor so the image fits
    # within BOTH bounds simultaneously, preserving aspect ratio. Never
    # exceeds 1.0, so a diagram already within bounds is never upscaled —
    # it keeps its true native size.
    scale = min(width_scale, height_scale, 1.0)

    final_width = native_width_in * scale
    final_height = native_height_in * scale

    return Inches(final_width), Inches(final_height)


# ─────────────────────────────────────────────────────────────────────────────
# Comment stripping
# ─────────────────────────────────────────────────────────────────────────────

def _strip_all_comments(doc: Document):
    """
    Remove ALL existing comments from a document.

    Final Fixed doc must be completely clean — zero comments.
    This strips:
    - Comments from previous manual reviewers (e.g. shobhit jaiswal)
    - Any comments carried over from the original document
    - Comment reference marks in paragraph runs

    Must be called immediately after copy_docx() in build_final_fixed_doc.
    """
    # Remove comment reference marks from all paragraphs
    for para in doc.paragraphs:
        for elem in para._element.xpath('.//*[local-name()="commentRangeStart"]'):
            elem.getparent().remove(elem)
        for elem in para._element.xpath('.//*[local-name()="commentRangeEnd"]'):
            elem.getparent().remove(elem)
        for elem in para._element.xpath('.//*[local-name()="commentReference"]'):
            elem.getparent().remove(elem)

    # Clear the comments XML part entirely
    try:
        part = doc.part
        if hasattr(part, '_comments_part') and part._comments_part is not None:
            comments_elem = part._comments_part._element
            if comments_elem is not None:
                for child in list(comments_elem):
                    comments_elem.remove(child)
    except Exception:
        pass

    # Also clear via relationship if accessible
    try:
        for rel in list(doc.part.rels.values()):
            if 'comments' in rel.reltype.lower():
                comments_part = rel.target_part
                root = comments_part._element
                for child in list(root):
                    root.remove(child)
                break
    except Exception:
        pass


# ─────────────────────────────────────────────────────────────────────────────
# Style normalization
# ─────────────────────────────────────────────────────────────────────────────
def normalize_heading_levels(doc):
    """
    IBM heading font-size standard:
      Heading 1 -> 12 pt
      Heading 2 -> 10 pt
      Heading 3 (and any deeper level) -> 10 pt
    All headings: Arial, Bold.

    Alignment rule: a heading that is ALREADY aligned (CENTER, RIGHT, or
    DISTRIBUTE) is left exactly as it is — never forced to LEFT. Only a
    heading with NO explicit alignment (i.e. broken/inherited/missing) is
    set to LEFT. This guarantees intentional title-page centering (e.g. a
    cover-page title) is never silently changed.
    """

    for para in doc.paragraphs:

        if not para.style:
            continue

        style_name = para.style.name.lower()

        if not style_name.startswith("heading"):
            continue

        try:
            level = int(style_name.replace("heading", "").strip())
        except Exception:
            continue

        target_size = HEADING_SIZE_BY_LEVEL.get(level, DEFAULT_HEADING_SIZE)

        # Preserve any already-intentional alignment; only fix missing/broken alignment.
        if para.alignment not in PRESERVED_ALIGNMENTS:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        for run in para.runs:
            run.font.name = "Arial"
            run.font.bold = True
            run.font.size = Pt(target_size)


def normalize_styles(doc: Document):
    """
    Fix only paragraphs whose font/size is inconsistent with the document's
    own dominant style. Does NOT force Arial — uses the document's own norm.
    Skips TOC entries entirely.
    Never touches heading paragraphs' size — normalize_heading_levels already
    enforces the IBM 12/10/10 standard and must remain the single source of
    truth for heading sizes.
    """
    dom_heading_font, dom_heading_size = _get_dominant_font(doc, is_heading=True)
    dom_body_font, dom_body_size = _get_dominant_font(doc, is_heading=False)

    # Update TOC style definitions to use dominant body font
    toc_style_names = [
        "toc 1", "toc 2", "toc 3", "toc 4", "toc 5",
        "toc 6", "toc 7", "toc 8", "toc 9", "TOC Heading"
    ]
    target_toc_font = dom_body_font or REQUIRED_BODY_FONT
    for sn in toc_style_names:
        try:
            style = doc.styles[sn]
            if style.font:
                style.font.name = target_toc_font
        except KeyError:
            pass

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else "Normal"
        style_lower = style_name.lower()

        # Never touch TOC paragraphs
        if style_lower.startswith("toc"):
            continue

        is_heading = style_lower.startswith("heading")
        target_font = dom_heading_font if is_heading else dom_body_font
        target_size = dom_heading_size if is_heading else dom_body_size

        if not target_font and not target_size:
            continue

        for run in para.runs:
            if not run.text.strip():
                continue

            run_font = run.font.name
            run_size = run.font.size.pt if run.font.size else None

            if run_font and target_font and run_font != target_font:
                run.font.name = target_font

            # Heading sizes are owned exclusively by normalize_heading_levels
            # (12/10/10 IBM standard) — never override them here.
            if is_heading:
                continue

            if run_size is not None and target_size is not None:
                # Never shrink a run that's LARGER than the dominant body
                # size. A run far above the body-text norm is, by
                # definition, intentional display/title text (e.g. a
                # cover-page title at 26pt while body text is 8.5pt) — not
                # an inconsistency to "fix". Only correct runs that are
                # close to, but not exactly matching, the dominant size
                # (e.g. an 8pt typo next to 8.5pt body text), or runs
                # smaller than the dominant size that look like accidental
                # shrinkage rather than intentional small print.
                if run_size > target_size:
                    continue
                if abs(run_size - target_size) > 0.5:
                    run.font.size = Pt(target_size)


def normalize_pseudo_heading_sizes(doc: Document) -> None:
    """
    Set every bold "pseudo-heading" paragraph to PSEUDO_HEADING_SIZE_PT
    (14pt), document-wide.

    A pseudo-heading is a paragraph that is NOT a real Word Heading style
    but visually functions as one: it has at least one bold run and is
    short (<= PSEUDO_HEADING_MAX_WORDS words). This is the exact same
    detection rule already used by justify_body_paragraphs() to decide
    which lines get LEFT alignment instead of JUSTIFY — reused here so
    sizing and alignment treat the same set of paragraphs consistently.

    This includes title-page lines (e.g. the cover title, "Theory
    Courseware", "Topics Covered") and any other bold standalone label
    throughout the body (e.g. "Unit 1 — Summary", a bolded glossary term
    lead-in). Real Heading 1/2/3 styled paragraphs are untouched here —
    they are governed exclusively by normalize_heading_levels()'s
    12/10/10pt rule. TOC entries are also skipped.

    Must run AFTER normalize_styles(), since normalize_styles() only
    skips shrinking runs already larger than the body-text norm — it does
    not raise smaller pseudo-headings up to 14pt. This function is the
    single source of truth for pseudo-heading size and runs last so its
    14pt value is never overwritten by an earlier step.
    """
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else "Normal"
        style_lower = style_name.lower()

        if style_lower.startswith("heading") or style_lower.startswith("toc"):
            continue

        text = para.text.strip()
        if not text:
            continue

        runs_with_text = [r for r in para.runs if r.text.strip()]
        has_any_bold = runs_with_text and any(r.bold for r in runs_with_text)
        word_count = len(re.sub(r'\s+', ' ', text).split())

        if has_any_bold and word_count <= PSEUDO_HEADING_MAX_WORDS:
            for run in para.runs:
                if run.text.strip():
                    run.font.size = Pt(PSEUDO_HEADING_SIZE_PT)


# ─────────────────────────────────────────────────────────────────────────────
# Spacing cleanup
# ─────────────────────────────────────────────────────────────────────────────

def clean_doc_spacing(doc: Document):
    """
    Clean extra spacing in paragraphs without destroying inline formatting.
    Edits run text in-place — never uses para.clear() which destroys bold/italic/links.
    Uses document's own dominant font, not hardcoded Arial.
    Skips TOC entries and image paragraphs.
    """
    dom_heading_font, dom_heading_size = _get_dominant_font(doc, is_heading=True)
    dom_body_font, dom_body_size = _get_dominant_font(doc, is_heading=False)

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else "Normal"
        style_lower = style_name.lower()

        # Never touch TOC entries, heading paragraphs, or image paragraphs.
        # Headings are formatted separately — applying clean_spacing() to them
        # can collapse intentional multi-word spacing and break alignment.
        if style_lower.startswith("toc") or style_lower.startswith("heading"):
            continue
        has_drawing = bool(para._element.xpath('.//*[local-name()="drawing"]'))
        if has_drawing:
            continue
        if not para.text or not para.text.strip():
            continue

        cleaned = clean_spacing(para.text)
        if cleaned == para.text.strip():
            continue

        target_font = dom_body_font
        target_size = dom_body_size

        meaningful_runs = [r for r in para.runs if r.text.strip()]

        if len(meaningful_runs) == 1:
            run = meaningful_runs[0]
            run.text = cleaned
            if target_font and run.font.name and run.font.name != target_font:
                run.font.name = target_font
            if target_size and run.font.size:
                if abs(run.font.size.pt - target_size) > 0.5:
                    run.font.size = Pt(target_size)

        elif len(meaningful_runs) > 1:
            # Only strip edges and fix double spaces within each run
            # Never collapse — preserves bold/italic/hyperlink formatting
            meaningful_runs[0].text = re.sub(r'  +', ' ', meaningful_runs[0].text.lstrip())
            meaningful_runs[-1].text = re.sub(r'  +', ' ', meaningful_runs[-1].text.rstrip())
            for run in meaningful_runs[1:-1]:
                run.text = re.sub(r'  +', ' ', run.text)
            first = meaningful_runs[0]
            if target_font and first.font.name and first.font.name != target_font:
                first.font.name = target_font
            if target_size and first.font.size:
                if abs(first.font.size.pt - target_size) > 0.5:
                    first.font.size = Pt(target_size)


# ─────────────────────────────────────────────────────────────────────────────
# Document helpers
# ─────────────────────────────────────────────────────────────────────────────

def copy_docx(input_path: Path) -> Document:
    return Document(str(input_path))


def _paragraph_has_meaningful_text(paragraph):
    return bool(paragraph.text and paragraph.text.strip())


def _get_best_paragraph(doc: Document, paragraph_index):
    if not doc.paragraphs:
        p = doc.add_paragraph(" ")
        return p, 0

    if paragraph_index is not None and 0 <= paragraph_index < len(doc.paragraphs):
        p = doc.paragraphs[paragraph_index]
        if _paragraph_has_meaningful_text(p) or p.runs:
            return p, paragraph_index

        for distance in range(1, 6):
            for idx in [paragraph_index - distance, paragraph_index + distance]:
                if 0 <= idx < len(doc.paragraphs):
                    candidate = doc.paragraphs[idx]
                    if _paragraph_has_meaningful_text(candidate) or candidate.runs:
                        return candidate, idx

    for idx, p in enumerate(doc.paragraphs):
        if _paragraph_has_meaningful_text(p) or p.runs:
            return p, idx

    return doc.paragraphs[0], 0


def _get_or_create_anchor_run(paragraph):
    for run in paragraph.runs:
        if run.text and run.text.strip():
            return run
    return paragraph.add_run(" ")


def _highlight_anchor_context(paragraph, max_runs=2):
    highlighted = 0
    for run in paragraph.runs:
        if run.text and run.text.strip():
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            highlighted += 1
            if highlighted >= max_runs:
                break
    if highlighted == 0 and paragraph.runs:
        paragraph.runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW


def insert_paragraph_after(paragraph, text=None, doc=None):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)

    if text:
        run = new_para.add_run(text)
        if doc is not None:
            dom_font, dom_size = _get_dominant_font(doc, is_heading=False)
            run.font.name = dom_font or REQUIRED_BODY_FONT
            run.font.size = Pt(dom_size or REQUIRED_BODY_SIZE_PT)
        else:
            run.font.name = REQUIRED_BODY_FONT
            run.font.size = Pt(REQUIRED_BODY_SIZE_PT)

    return new_para


def _remove_paragraph(paragraph):
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


# ─────────────────────────────────────────────────────────────────────────────
# Issue tag mapping
# ─────────────────────────────────────────────────────────────────────────────

_ISSUE_TYPE_TAG = {
    "grammar":                          "[Grammar]",
    "spacing":                          "[Format]",
    "formatting_body":                  "[Format]",
    "formatting_heading":               "[Format]",
    "heading_hierarchy":                "[Structure]",
    "navigation_pane":                  "[Structure]",
    "repeated_heading":                 "[Structure]",
    "duplicate_paragraph":              "[Duplicate]",
    "near_duplicate_paragraph":         "[Duplicate]",
    "semantic_duplicate_paragraph":     "[Duplicate]",
    "possible_duplicate_topic":         "[Duplicate]",
    "image_caption_and_source_missing": "[Image]",
    "image_caption_missing":            "[Image]",
    "image_source_missing":             "[Image]",
    "figure_number_missing":            "[Image]",
    "ocr_low_confidence":               "[OCR]",
    "content_accuracy":                 "[Accuracy]",
    "quiz_missing":                     "[Quiz]",
    "diagram_recommended":              "[Diagram]",
    "diagram_generated":                "[Diagram]",
    "diagram_generation_failed":        "[Diagram]",
    "advanced_visual_generated":        "[Visual]",
    "possible_ibm_similarity":          "[Plagiarism]",
    "possible_open_source_similarity":  "[Plagiarism]",
    "winston_scan_error":               "[Plagiarism]",
}


def _get_issue_tag(issue_type: str) -> str:
    return _ISSUE_TYPE_TAG.get(
        issue_type,
        f"[{issue_type.replace('_', ' ').title()}]"
    )


def _safe_comment_text(issue):
    tag = _get_issue_tag(issue.issue_type)
    return (
        f"{tag}\n"
        f"Severity: {issue.severity.upper()}\n"
        f"Location: {issue.location}\n"
        f"Comment: {issue.message}\n"
        f"Suggested Fix: {issue.suggested_fix}"
    )


# ─────────────────────────────────────────────────────────────────────────────
# Review Comments document builders
# ─────────────────────────────────────────────────────────────────────────────

def add_inline_comments_to_doc(doc: Document, state):
    """Add all pipeline review issues as Word comments in the Review doc."""
    if not doc.paragraphs:
        doc.add_paragraph(" ")

    for issue in state.issues:
        paragraph, _ = _get_best_paragraph(doc, issue.paragraph_index)
        anchor_run = _get_or_create_anchor_run(paragraph)
        _highlight_anchor_context(paragraph)

        try:
            doc.add_comment(
                anchor_run,
                text=_safe_comment_text(issue),
                author=COMMENT_AUTHOR,
                initials=COMMENT_INITIALS,
            )
        except Exception:
            fallback_run = paragraph.add_run(" ")
            fallback_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            try:
                doc.add_comment(
                    fallback_run,
                    text=_safe_comment_text(issue),
                    author=COMMENT_AUTHOR,
                    initials=COMMENT_INITIALS,
                )
            except Exception:
                pass


def add_ocr_summary_to_review_doc(doc: Document, state):
    # Internal pipeline info — not shown in Review Comments document
    pass


def _insert_review_summary(doc: Document, state) -> None:
    """
    Inserts a summary page at the very beginning of the Review Comments doc.
    Shows total issues, severity breakdown, top issue types, and units detected.
    """
    issues = getattr(state, "issues", [])
    units  = getattr(state, "units", [])

    severity_counts = Counter(i.severity for i in issues)
    type_counts     = Counter(i.issue_type for i in issues)

    body = doc.element.body

    def _prepend_para(text: str, bold: bool = False,
                      size_pt: int = None, space_after: int = 4) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        if size_pt:
            run.font.size = Pt(size_pt)
        pPr = p._p.get_or_add_pPr()
        spc = OxmlElement("w:spacing")
        spc.set(qn("w:after"), str(space_after * 20))
        pPr.append(spc)
        body.remove(p._p)
        body.insert(0, p._p)

    # Build reverse so final doc reads top-to-bottom

    # Page break after summary
    p_break = doc.add_paragraph()
    p_break.add_run().add_break(WD_BREAK.PAGE)
    body.remove(p_break._p)
    body.insert(0, p_break._p)

    # Units
    if units:
        for unit in reversed(units):
            quiz_note = ""
            if unit.quiz_present:
                quiz_note = "  ✓ quiz present"
            elif unit.generated_quiz:
                quiz_note = f"  ({len(unit.generated_quiz)} questions generated)"
            _prepend_para(f"  • {unit.title}{quiz_note}")
        _prepend_para("Units / Chapters Detected", bold=True)

    # Top issue types
    if type_counts:
        for issue_type, count in reversed(type_counts.most_common(8)):
            tag = _get_issue_tag(issue_type)
            _prepend_para(f"  {tag}  —  {count} issue(s)")
        _prepend_para("Top Issue Types", bold=True)

    # Severity breakdown
    for label in reversed(["info", "low", "medium", "high"]):
        count = severity_counts.get(label, 0)
        icon  = {"high": "🔴", "medium": "🟠", "low": "🟡", "info": "🔵"}.get(label, "⚪")
        _prepend_para(f"  {icon}  {label.upper()}: {count}")
    _prepend_para("Issues by Severity", bold=True)

    # Total
    _prepend_para(f"Total Issues Found: {len(issues)}", bold=True, size_pt=12)

    # Title
    _prepend_para("COURSEWARE REVIEW SUMMARY", bold=True, size_pt=16, space_after=8)

def _inject_image_issues(state, original_file: Path = None) -> None:
    """
    Pre-populate state.issues with ReviewIssue entries for every image
    that is missing a caption or source link.
    Must be called BEFORE build_review_comments_doc so the review doc
    reflects every auto-insertion that build_final_fixed_doc will make.
    """  
    from image_source_finder import find_sources_for_images
    # Cache on state so build_final_fixed_doc reuses the exact same resolved URLs —
    # calling find_sources_for_images twice can produce different results (API timing,
    # network variance) which causes Review doc and Fixed doc to diverge.
    if not getattr(state, '_cached_image_sources', None):
        state._cached_image_sources = find_sources_for_images(state)
    image_sources = state._cached_image_sources

    seen_caption = set()
    seen_source  = set()
    if original_file is not None:
        _tmp_doc = copy_docx(original_file)
        figure_counter = _get_next_figure_number(_tmp_doc)
        del _tmp_doc
    else:
        figure_counter = 1

    existing_issues = {
    (i.issue_type, i.paragraph_index)
    for i in getattr(state, "issues", [])
    }
    
    for img in state.images:
        if img.paragraph_index is None:
            continue
        target_idx = (
            img.nearest_caption_paragraph_index
            if img.nearest_caption_paragraph_index is not None
            else img.paragraph_index
        )
        source_url = image_sources.get(img.filename)
        needs_caption = (
            not img.nearest_caption_text
            and target_idx not in seen_caption
        )

        needs_source = (
            not img.has_source_link
            and source_url
            and "Add source URL/reference here" not in source_url
            and (target_idx, source_url) not in seen_source
        )
        if not (needs_caption or needs_source):
            continue

        msg_parts = []
        fix_parts = []

        if needs_caption:
            msg_parts.append(
                f"Image source reference added: '{source_url}'")
            fix_parts.append("Replace the placeholder caption with an accurate figure description.")
            seen_caption.add(target_idx)
            figure_counter += 1

        if needs_source:
            msg_parts.append(
                "Auto-inserted source URL placeholder: "
                "'Source: [Add source URL/reference here]'"
            )
            fix_parts.append(
                "Verify that the detected source reference is correct."
            )
            seen_source.add((target_idx, source_url))

        if needs_caption and needs_source:
            issue_type = "image_caption_and_source_missing"
        elif needs_caption:
            issue_type = "image_caption_missing"
        else:
            issue_type = "image_source_missing"

        issue_key = (issue_type, target_idx)

        if issue_key not in existing_issues:
            state.issues.append(
                ReviewIssue(
                    issue_type=issue_type,
                    severity="medium",
                    paragraph_index=target_idx,
                    location=f"Image: {img.filename or 'unknown'} (paragraph {target_idx})",
                    message=" | ".join(msg_parts),
                    suggested_fix=" ".join(fix_parts),
                )
            )

def build_review_comments_doc(original_file: Path, state, output_path: Path):
    """
    Build the Review Comments document.
    Copies the original, adds inline Word comments for every issue,
    and prepends a summary page.
    Does NOT strip existing comments — reviewer may want to see them.

    Also inserts the same visible figure-caption and source-link placeholder
    lines that build_final_fixed_doc inserts, so both output documents stay
    in alignment (same placeholders, same positions, same figure numbers).
    """
    doc = copy_docx(original_file)
    add_inline_comments_to_doc(doc, state)

    # ── Insert visible caption / source placeholders ──────────────────────
    # Mirrors build_final_fixed_doc Step 7 exactly so the Review doc shows
    # the same "Figure N: ..." and "Source: ..." lines that the Fixed doc has.
    # Uses the image_sources dict already cached by _inject_image_issues so
    # both documents always reference the same resolved URLs.
    from image_source_finder import find_sources_for_images
    if not getattr(state, '_cached_image_sources', None):
        state._cached_image_sources = find_sources_for_images(state)
    image_sources = state._cached_image_sources

    figure_counter = _get_next_figure_number(doc)
    caption_inserted_at = set()
    source_inserted_at = set()

    for img in state.images:
        if img.paragraph_index is None:
            continue
        target_idx = (
            img.nearest_caption_paragraph_index
            if img.nearest_caption_paragraph_index is not None
            else img.paragraph_index
        )
        _img_source_url = image_sources.get(img.filename)
        needs_caption = (not img.nearest_caption_text) and (target_idx not in caption_inserted_at)
        needs_source = (
                (not img.has_source_link)
                and _img_source_url
                and "Add source URL/reference here" not in _img_source_url
                and ((target_idx, _img_source_url) not in source_inserted_at)
            )      
         
        if needs_caption or needs_source:
            figure_text = (
                    
                    f"Figure {figure_counter}"
                    if needs_caption else None
                )
            source_text = (
                
                f"{_img_source_url}"
                if needs_source else None
            )

            _append_caption_then_source(doc, target_idx, figure_text, source_text)

            if needs_caption:
                caption_inserted_at.add(target_idx)
                figure_counter += 1
            if needs_source:
                source_inserted_at.add((target_idx, _img_source_url))
    # ── End caption / source insertion ───────────────────────────────────

    add_ocr_summary_to_review_doc(doc, state)
    _insert_review_summary(doc, state)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


# ─────────────────────────────────────────────────────────────────────────────
# Final Fixed document builders
# ─────────────────────────────────────────────────────────────────────────────

def _is_figure_caption(text: str) -> bool:
    if not text:
        return False
    return bool(re.search(r'^\s*(figure|fig\.?|fig)\s*\d+', text, flags=re.IGNORECASE))


def _get_next_figure_number(doc: Document) -> int:
    max_fig = 0
    for para in doc.paragraphs:
        m = re.search(r'figure\s*(\d+)', para.text, flags=re.IGNORECASE)
        if m:
            num = int(m.group(1))
            if num > max_fig:
                max_fig = num
    return max_fig + 1


def _is_list_item(para) -> bool:
    """Return True if the paragraph is part of a numbered or bulleted list."""
    return bool(para._element.xpath('.//*[local-name()="numPr"]'))


def _append_caption_then_source(doc: Document, idx: int, figure_text: str, source_text: str):
    """Insert figure caption then source link after an image paragraph.
    
    If the image sits inside a numbered list (e.g. between step 3 and step 4),
    we walk forward past all consecutive list items before inserting so the
    list is never split in the middle.
    """
    if idx is None or idx < 0 or idx >= len(doc.paragraphs):
        anchor = doc.paragraphs[-1] if doc.paragraphs else doc.add_paragraph()
    else:
        anchor = doc.paragraphs[idx]

    # Walk forward past any immediately following list-item paragraphs so we
    # don't insert captions in the middle of a numbered/bulleted list.
    current_idx = idx if idx is not None else 0
    while current_idx + 1 < len(doc.paragraphs):
        next_para = doc.paragraphs[current_idx + 1]
        # Stop walking if the next paragraph is not a list item, is an image,
        # or already has a figure caption (don't skip past the end of the list's section).
        if not _is_list_item(next_para):
            break
        has_drawing = bool(next_para._element.xpath('.//*[local-name()="drawing"]'))
        if has_drawing:
            break
        current_idx += 1
        anchor = next_para

    def _has_caption_nearby(para_idx):
        for check_idx in range(para_idx, min(para_idx + 4, len(doc.paragraphs))):
            if _is_figure_caption(doc.paragraphs[check_idx].text):
                return True
        return False

    caption_anchor = anchor

    if figure_text and not _has_caption_nearby(current_idx):
        new_cap = insert_paragraph_after(caption_anchor)
        run = new_cap.add_run(figure_text)
        run.italic = True
        caption_anchor = new_cap

    if source_text:
        new_src = insert_paragraph_after(caption_anchor)
        run = new_src.add_run(source_text)
        run.italic = True


def add_missing_summary_headings(doc: Document):
    """
    Detect paragraphs that begin with summary/conclusion phrases and
    insert a Heading 3 before them if one is not already present.
    Only triggers when the paragraph clearly represents a standalone summary
    section — not mid-paragraph phrases in technical content.
    """
    SUMMARY_PATTERNS = [
        r'^in summary\s*[,:]',
        r'^to summarize\s*[,:]',
        r'^in conclusion\s*[,:]',
        r'^to conclude\s*[,:]',
        r'^summary\s*:',
        r'^conclusion\s*:',
        r'^overall\s*[,:]',
        r'^in brief\s*[,:]',
    ]

    dom_heading_font, dom_heading_size = _get_dominant_font(doc, is_heading=True)
    h_font = dom_heading_font or REQUIRED_HEADING_FONT
    h_size = dom_heading_size or REQUIRED_HEADING_SIZE_PT

    to_insert = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().lower()
        style = (para.style.name if para.style else "Normal").lower()
        if style.startswith("heading") or style.startswith("toc"):
            continue
        if any(re.match(p, text) for p in SUMMARY_PATTERNS):
            # Long paragraphs are flowing prose, not standalone summary sections
            if len(para.text.split()) > 30:
                continue
            # Skip paragraphs with technical content (cron, variables, symbols)
            if re.search(r'\d+\s+[\d*]|\$[\w?]|[{}();]', para.text):
                continue
            if i > 0:
                prev_style = (doc.paragraphs[i - 1].style.name if doc.paragraphs[i - 1].style else "").lower()
                if prev_style.startswith("heading"):
                    continue
            to_insert.append(i)

    for i in sorted(to_insert, reverse=True):
        anchor = doc.paragraphs[i]
        new_p = OxmlElement("w:p")
        anchor._p.addprevious(new_p)
        new_para = Paragraph(new_p, anchor._parent)
        try:
            new_para.style = doc.styles["Heading 3"]
        except KeyError:
            pass
        run = new_para.add_run("Summary")
        run.font.name = h_font
        run.font.size = Pt(h_size)
        run.bold = True


def remove_exact_duplicate_body_paragraphs(doc: Document):
    """
    Remove duplicate body paragraphs that appear within MAX_DUPLICATE_DISTANCE
    paragraphs of each other. Distant duplicates are likely intentional repetition.
    Never removes headings, TOC entries, or image paragraphs.
    """
    MAX_DUPLICATE_DISTANCE = 20
    seen = {}
    to_remove = []

    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else "Normal"
        style_lower = style_name.lower()
        if style_lower.startswith("heading") or style_lower.startswith("toc"):
            continue
        has_drawing = bool(para._element.xpath('.//*[local-name()="drawing"]'))
        if has_drawing:
            continue
        norm = normalize_text(para.text).lower()
        if not norm or len(norm) < 80:
            continue
        # Never remove source lines or figure captions — multiple images may
        # legitimately share the same portal URL or a generic caption placeholder.
        if re.match(r'^\s*(source\s*:|figure\s*\d+)', norm, re.IGNORECASE):
            continue
        if norm in seen:
            first_idx = seen[norm]
            if i - first_idx <= MAX_DUPLICATE_DISTANCE:
                to_remove.append(para)
            seen[norm] = i
        else:
            seen[norm] = i

    for para in to_remove:
        _remove_paragraph(para)


def apply_rewrites_to_doc(doc: Document, state):
    """
    Apply LLM rewrite suggestions in-place.
    Edits first run text only — preserves all inline formatting.
    Uses document's own dominant font, not hardcoded Arial.
    """
    dom_heading_font, dom_heading_size = _get_dominant_font(doc, is_heading=True)
    dom_body_font, dom_body_size = _get_dominant_font(doc, is_heading=False)

    for item in getattr(state, "rewrite_suggestions", []):
        if item.skipped or not item.applied:
            continue
        idx = item.paragraph_index
        if idx is None or idx < 0 or idx >= len(doc.paragraphs):
            continue

        para = doc.paragraphs[idx]
        has_image = para._element.xpath(
            './/*[local-name()="drawing"] | .//*[local-name()="pict"]'
        )
        if has_image:
            continue

        style_name = (para.style.name if para.style else "Normal").lower()
        is_heading = style_name.startswith("heading")
        target_font = dom_heading_font if is_heading else dom_body_font
        target_size = dom_heading_size if is_heading else dom_body_size

        meaningful_runs = [r for r in para.runs if r.text.strip()]

        if not meaningful_runs:
            run = para.add_run(item.rewritten_text)
            if target_font:
                run.font.name = target_font
            if target_size:
                run.font.size = Pt(target_size)
            if is_heading:
                run.bold = True
        else:
            meaningful_runs[0].text = item.rewritten_text
            for run in meaningful_runs[1:]:
                run.text = ""
            first_run = meaningful_runs[0]
            if target_font and first_run.font.name and first_run.font.name != target_font:
                first_run.font.name = target_font
            if target_size and first_run.font.size:
                if abs(first_run.font.size.pt - target_size) > 0.5:
                    first_run.font.size = Pt(target_size)
            if is_heading:
                first_run.bold = True


def _find_unit_end_in_doc(doc: Document, unit_title: str) -> Optional[int]:
    """
    Find the paragraph index at which a unit's content ends in the live
    document, so a generated quiz can be inserted immediately after it —
    i.e. at the true END of the unit, never near the start.

    Drift-safe: re-locates the heading by text match in the live document
    instead of trusting a pre-recorded index, since earlier pipeline steps
    (rewrites, caption insertion, duplicate removal) can shift indices.
    """
    unit_title_norm = unit_title.strip().lower()
    unit_heading_idx = None
    unit_level = None

    # A bold short line, or a same-level Word heading, only counts as a
    # NEXT-unit boundary if it actually looks like a chapter/unit/module/
    # lesson title (e.g. "Unit 2: ...") AND carries a DIFFERENT number than
    # the current unit. Without the number check, in-unit labels like
    # "Unit 1 — Summary" would be mistaken for the start of the next unit.
    # Without the keyword check, any incidental bold line ("Topics
    # Covered", a key term, a callout label) would be mistaken for a
    # boundary too — both truncate the unit early and place the quiz near
    # the top of the document instead of at the true end.
    CHAPTER_TITLE_PATTERN = re.compile(
        r'^(chapter|unit|module|lesson)\s+(\d+|[ivxIVX]+)\b', re.IGNORECASE
    )
    current_unit_number_match = re.search(r'\b(\d+|[ivxIVX]+)\b', unit_title)
    current_unit_number = current_unit_number_match.group(1).lower() if current_unit_number_match else None

    def _is_next_unit_title(text: str) -> bool:
        m = CHAPTER_TITLE_PATTERN.match(text)
        if not m:
            return False
        candidate_number = m.group(2).lower()
        # Same number as the current unit -> it's an in-unit label
        # ("Unit 1 — Summary"), not the next unit.
        return candidate_number != current_unit_number

    # Find the heading text in the live document. Only match real
    # Word-styled headings (Heading 1/2/3...) — never bold pseudo-headings —
    # so an earlier cover-page repeat of the same title text isn't mistaken
    # for the real section heading.
    for i, para in enumerate(doc.paragraphs):
        style_name = (para.style.name if para.style else "").lower()
        is_word_heading = style_name.startswith("heading")
        if not is_word_heading:
            continue
        if para.text.strip().lower() == unit_title_norm:
            unit_heading_idx = i
            parts = style_name.split()
            if len(parts) >= 2 and parts[1].isdigit():
                unit_level = int(parts[1])
            break

    if unit_heading_idx is None:
        print(f"[QUIZ INSERT] Could not find heading '{unit_title}' in live doc — using end_paragraph_index fallback")
        return None

    for i in range(unit_heading_idx + 1, len(doc.paragraphs)):
        para_i = doc.paragraphs[i]
        style_name = (para_i.style.name if para_i.style else "").lower()
        is_word_heading = style_name.startswith("heading")
        text_i = para_i.text.strip()
        is_bold_heading = (
            bool(text_i)
            and len(text_i.split()) <= 12
            and bool(para_i.runs)
            and all(r.bold for r in para_i.runs if r.text.strip())
        )
        if is_word_heading:
            parts = style_name.split()
            if len(parts) >= 2 and parts[1].isdigit():
                level = int(parts[1])
                # Same/shallower heading level alone isn't reliable — some
                # documents use Heading 1 for both the unit title AND its
                # numbered subsections (1.1, 1.2, 1.3...). Only treat it as
                # the start of the NEXT unit if it's strictly shallower, or
                # it's the same level but actually reads like a new,
                # DIFFERENT-numbered chapter/unit title — not a subsection
                # heading like "1.1 Introduction..." and not an in-unit
                # label like "Unit 1 — Summary".
                if unit_level is None or level < unit_level:
                    return i - 1
                if level == unit_level and _is_next_unit_title(text_i):
                    return i - 1
        elif is_bold_heading and _is_next_unit_title(text_i):
            return i - 1

    return len(doc.paragraphs) - 1


def insert_generated_quizzes_after_units(doc: Document, state):
    """
    Insert exactly one generated quiz at the end of each unit that is
    missing one.

    - If a unit already has a quiz (unit.quiz_present is True), nothing is
      inserted for that unit — no duplicate is ever created.
    - If a unit has no quiz, exactly one quiz is generated and inserted at
      the TRUE end of that unit (see _find_unit_end_in_doc), never near
      the start.

    DRIFT-SAFE: Uses heading text match to find position in the live document
    instead of the pre-recorded end_paragraph_index which drifts after rewrites
    and caption insertions.

    Processing order: bottom-to-top so earlier insertions don't shift later positions.
    Font: uses document's own dominant font, not hardcoded Arial.
    """
    units = [
        u for u in getattr(state, "units", [])
        if (not u.quiz_present) and u.generated_quiz
    ]
    if not units:
        return

    # Resolve all positions before any insertion
    resolved = []
    for unit in units:
        insert_idx = _find_unit_end_in_doc(doc, unit.title)
        if insert_idx is None:
            insert_idx = unit.end_paragraph_index  # fallback
        if insert_idx is None:
            continue
        insert_idx = max(0, min(insert_idx, len(doc.paragraphs) - 1))
        resolved.append((unit, insert_idx))

    # Sort bottom-to-top
    resolved.sort(key=lambda x: x[1], reverse=True)

    # Cache fonts once — expensive to compute per unit on large docs
    dom_heading_font, dom_heading_size = _get_dominant_font(doc, is_heading=True)
    dom_body_font, dom_body_size = _get_dominant_font(doc, is_heading=False)
    h_font = dom_heading_font or REQUIRED_HEADING_FONT
    h_size = dom_heading_size or REQUIRED_HEADING_SIZE_PT
    b_font = dom_body_font or REQUIRED_BODY_FONT
    b_size = dom_body_size or REQUIRED_BODY_SIZE_PT

    for unit, insert_idx in resolved:
        anchor_para = doc.paragraphs[insert_idx]

        # Page break before quiz
        page_break_para = insert_paragraph_after(anchor_para)
        page_break_para.add_run().add_break(WD_BREAK.PAGE)

        # Quiz heading
        current_anchor = insert_paragraph_after(page_break_para, f"Quiz – {unit.title}")
        current_anchor.style = doc.styles["Normal"]
        current_anchor.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if current_anchor.runs:
            current_anchor.runs[0].bold = True
            current_anchor.runs[0].font.name = h_font
            current_anchor.runs[0].font.size = Pt(h_size)

        current_type = None

        for item in unit.generated_quiz:
            # Section subheading per question type
            if item.qtype != current_type:
                current_type = item.qtype
                current_anchor = insert_paragraph_after(current_anchor, current_type)
                current_anchor.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if current_anchor.runs:
                    current_anchor.runs[0].bold = True
                    current_anchor.runs[0].font.name = b_font
                    current_anchor.runs[0].font.size = Pt(b_size)

            # Question
            current_anchor = insert_paragraph_after(
                current_anchor,
                f"Q. {item.question}"
            )

            current_anchor.style = doc.styles["Normal"]
            current_anchor.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if current_anchor.runs:
                current_anchor.runs[0].font.name = b_font
                current_anchor.runs[0].font.size = Pt(b_size)

            # MCQ options
            if item.options:
                for opt in item.options:
                    current_anchor = insert_paragraph_after(current_anchor, f"- {opt}")
                    current_anchor.style = doc.styles["Normal"]
                    current_anchor.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    if current_anchor.runs:
                        current_anchor.runs[0].font.name = b_font
                        current_anchor.runs[0].font.size = Pt(b_size)
            
            # Answer
            if item.answer:
                current_anchor = insert_paragraph_after(current_anchor, f"Answer: {item.answer}")
                current_anchor.style = doc.styles["Normal"]
                current_anchor.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if current_anchor.runs:
                    current_anchor.runs[0].font.name = b_font
                    current_anchor.runs[0].font.size = Pt(b_size)

def insert_generated_diagrams(doc: Document, state):
    """
    Insert auto-generated diagrams into the document.
    Images are sized to fit within both the page body width AND a maximum
    height (MAX_IMAGE_HEIGHT_INCHES), preserving native aspect ratio, then
    centred. This prevents a naturally tall/narrow diagram from being
    stretched into a multi-page-tall banner image.
    """
    diagrams = sorted(
        getattr(state, "generated_diagrams", []),
        key=lambda x: x["paragraph_index"],
        reverse=True
    )
    for item in diagrams:
        img_path = item.get("image_path")
        para_idx = item.get("paragraph_index")
        if not img_path or para_idx is None:
            continue
        if para_idx < 0 or para_idx >= len(doc.paragraphs):
            continue
        anchor_para = doc.paragraphs[para_idx]
        caption_para = insert_paragraph_after(
            anchor_para,
            f"[Auto-generated {item.get('diagram_type', 'diagram')} diagram]"
        )
        if caption_para.runs:
            caption_para.runs[0].italic = True
            caption_para.runs[0].font.name = REQUIRED_BODY_FONT
            caption_para.runs[0].font.size = Pt(REQUIRED_BODY_SIZE_PT)
        image_para = insert_paragraph_after(caption_para)
        image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        safe_width, safe_height = _safe_image_dimensions(doc, img_path, DIAGRAM_IMAGE_WIDTH_INCHES)
        if safe_height is not None:
            image_para.add_run().add_picture(img_path, width=safe_width, height=safe_height)
        else:
            image_para.add_run().add_picture(img_path, width=safe_width)


def insert_advanced_visuals(doc: Document, state):
    """
    Insert advanced generated visuals into the document.
    Images are sized to fit within both the page body width AND a maximum
    height (MAX_IMAGE_HEIGHT_INCHES), preserving native aspect ratio, then
    centred. This prevents a naturally tall/narrow diagram from being
    stretched into a multi-page-tall banner image.
    """
    visuals = sorted(
        [v for v in getattr(state, "generated_visuals", []) if v.status == "success" and v.image_path],
        key=lambda x: x.paragraph_index,
        reverse=True,
    )
    for item in visuals:
        idx = item.paragraph_index
        if idx is None or idx < 0 or idx >= len(doc.paragraphs):
            continue
        anchor_para = doc.paragraphs[idx]
        label_para = insert_paragraph_after(
            anchor_para,
            f"[Auto-generated {item.visual_type.replace('_', ' ')} visual: {item.title}]"
        )
        if label_para.runs:
            label_para.runs[0].italic = True
            label_para.runs[0].font.name = REQUIRED_BODY_FONT
            label_para.runs[0].font.size = Pt(REQUIRED_BODY_SIZE_PT)
        image_para = insert_paragraph_after(label_para)
        image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        safe_width, safe_height = _safe_image_dimensions(doc, item.image_path, ADVANCED_VISUAL_WIDTH_INCHES)
        if safe_height is not None:
            image_para.add_run().add_picture(item.image_path, width=safe_width, height=safe_height)
        else:
            image_para.add_run().add_picture(item.image_path, width=safe_width)


def polish_doc_structure(doc: Document):
    """Remove consecutive blank paragraphs and empty heading-styled paragraphs."""
    to_remove = []
    prev_blank = False

    for para in doc.paragraphs:
        has_drawing = bool(para._element.xpath('.//*[local-name()="drawing"]'))
        style_lower = (para.style.name if para.style else "").lower()
        is_heading_style = style_lower.startswith("heading")
        is_blank_text = not (para.text and para.text.strip())

        # Remove heading-styled paragraphs with no visible text
        # (these show up as empty entries in the Navigation Pane)
        if is_heading_style and is_blank_text and not has_drawing:
            to_remove.append(para)
            prev_blank = False
            continue

        # Remove consecutive blank non-heading paragraphs
        is_blank = is_blank_text and not has_drawing
        if is_blank and prev_blank:
            to_remove.append(para)
        prev_blank = is_blank

    for para in to_remove:
        _remove_paragraph(para)

def _is_quiz_paragraph(text: str) -> bool:
    if not text:
        return False

    text = text.strip()

    patterns = [
        r'^quiz\b',
        r'^multiple choice',
        r'^fill in the blank',
        r'^true or false',
        r'^two[- ]mark',
        r'^four[- ]mark',

        r'^q[\.\)]',
        r'^q\d+[\.\)]',

        r'^\d+[\.\)]\s',
        r'^[A-D][\.\)]\s',

        r'^answer\s*:',
        r'^reason\s*:',
    ]

    return any(
        re.match(pattern, text, re.IGNORECASE)
        for pattern in patterns
    )

def justify_body_paragraphs(doc: Document):
    """
    Apply justified alignment only to substantial body paragraphs.
    Skips headings, images, TOC, centred/right/distributed text, short
    lines, and list-style content.

    Alignment rule (applies everywhere in this function): a paragraph that
    is ALREADY intentionally aligned — CENTER, RIGHT, or DISTRIBUTE — is
    NEVER changed. Only paragraphs with broken/missing/incorrect alignment
    (i.e. anything else, including stray JUSTIFY on a heading or list line)
    get forced to LEFT. This guarantees pre-existing centered titles, right-
    aligned captions, etc. are always preserved exactly as the author set
    them, while genuinely broken alignment is repaired.
    """
    LIST_PATTERN = re.compile(
        r'^\s*('
        r'\d+[\.\)]\s'                      # 1. or 1)
        r'|\d+\.[A-Za-z]'                   # 1.If... — inline quiz, no space after dot
        r'|[A-Da-d][\.\)]\s'                # a. or a)
        r'|[A-D][\.\)]\s'                   # A) or A. — single MCQ option line (e.g. "A) 0")
        r'|[-•·▪▸]\s'                       # bullet symbols
        r'|Step\s+\d+\s*[:\-–]'            # Step 1: or Step 1 -
        r'|Note\s*[:\-–]'                   # Note: or Note -
        r'|Warning\s*[:\-–]'               # Warning:
        r'|Example\s*[:\-–]'               # Example:
        r'|Tip\s*[:\-–]'                   # Tip:
        r'|Important\s*[:\-–]'             # Important:
        r'|Answer\s*[:\-–]'                # Answer: — quiz answer lines
        r'|Reason\s*[:\-–]'                # Reason: — quiz reason lines
        r')',
        re.IGNORECASE
    )

    # Pattern to detect quiz question lines (e.g. "Q. ...", "Q1. ...", "Q1) ...")
    # These must keep LEFT alignment, not be justified.
    QUIZ_QUESTION_PATTERN = re.compile(
        r'^\s*Q[\.\)]\s'           # Q. or Q)
        r'|^\s*Q\d+[\.\)]\s',      # Q1. or Q1)
        re.IGNORECASE
    )

    # Pattern to detect quiz section subheadings (e.g. "Multiple Choice Questions",
    # "Fill in the Blanks", "True or False", "Short Answer Questions")
    QUIZ_SECTION_PATTERN = re.compile(
        r'^\s*(multiple\s+choice|fill\s+in\s+the\s+blank|true\s+or\s+false|short\s+answer'
        r'|long\s+answer|descriptive\s+question|match\s+the\s+following|quiz\b)',
        re.IGNORECASE
    )

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else "Normal"
        style_lower = style_name.lower()

        # ── Headings (Word heading style) and TOC ───────────────────────────
        # Preserve any already-intentional alignment (CENTER/RIGHT/DISTRIBUTE).
        # Only force LEFT when alignment is missing or genuinely broken
        # (e.g. accidentally JUSTIFY).
        if style_lower.startswith("heading") or style_lower.startswith("toc"):
            if para.alignment not in PRESERVED_ALIGNMENTS:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue

        # Skip image paragraphs
        has_drawing = bool(para._element.xpath('.//*[local-name()="drawing"]'))
        if has_drawing:
            continue

        text_check = para.text.strip()

        # ── Pseudo-headings: bold paragraphs not using a Word heading style ─
        # Use ANY-bold check (not all-bold) because heading runs can be mixed.
        # Use <= 20 word limit to catch multi-word headings like
        # "The Principle of Localization:" that have extra spaces in source doc.
        # Same preservation rule as real headings: don't touch an existing
        # intentional CENTER/RIGHT/DISTRIBUTE alignment.
        if text_check:
            runs_with_text = [r for r in para.runs if r.text.strip()]
            has_any_bold = runs_with_text and any(r.bold for r in runs_with_text)
            # Collapse internal spaces before word-counting so "Fourier    Conjugates:"
            # isn't counted as many words due to spacing artifacts in source doc.
            word_count = len(re.sub(r'\s+', ' ', text_check).split())
            if has_any_bold and word_count <= 20:
                if para.alignment not in PRESERVED_ALIGNMENTS:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                continue

        # Skip paragraphs with explicit CENTER/RIGHT/DISTRIBUTE alignment —
        # never touch intentional alignment anywhere in this function.
        if para.alignment in PRESERVED_ALIGNMENTS:
            continue

        text = text_check

        if _is_quiz_paragraph(text):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue
        # Skip short lines — justifying these causes ugly extra spacing
        if len(text.split()) < 8:
            continue

        # Quiz/list lines: actively reset to LEFT (not just skip).
        # Skipping alone leaves any pre-existing JUSTIFY from the original doc intact.
        # Explicit LEFT ensures the original quiz alignment is always restored.
        if (
            LIST_PATTERN.match(text)
            or _INLINE_MCQ_PATTERN.search(text)
            or QUIZ_QUESTION_PATTERN.match(text)
            or QUIZ_SECTION_PATTERN.match(text)
        ):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue

        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def force_quiz_alignment(doc):
    """
    Final safety pass.
    Forces quiz content to LEFT alignment even if a previous
    formatting step accidentally justified it. Quiz structural lines
    (questions, options, answers) are never legitimately CENTER/RIGHT in
    this document type, so this targeted override is intentional and does
    not violate the general "never touch intentional alignment" rule used
    elsewhere — it only matches narrow, structurally-identifiable quiz syntax.
    """

    QUIZ_SECTION_PATTERN = re.compile(
        r'^\s*(multiple\s+choice|fill\s+in\s+the\s+blank|true\s+or\s+false|'
        r'two[- ]mark|four[- ]mark|short\s+answer|long\s+answer|quiz\b)',
        re.IGNORECASE
    )

    QUIZ_QUESTION_PATTERN = re.compile(
        r'^\s*Q[\.\)]\s'
        r'|^\s*Q\d+[\.\)]\s'
        r'|^\s*\d+[\.\)]\s'
        r'|^\s*\d+\.[A-Za-z]'
        r'|^\s*Question\s+\d+',
        re.IGNORECASE
    )

    for para in doc.paragraphs:
        text = (para.text or "").strip()

        if (
            _is_quiz_paragraph(text)
            or QUIZ_SECTION_PATTERN.match(text)
            or QUIZ_QUESTION_PATTERN.match(text)
            or _is_list_item(para)
        ):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

# ─────────────────────────────────────────────────────────────────────────────
# Table of Contents
# ─────────────────────────────────────────────────────────────────────────────

def _document_has_toc(doc: Document) -> bool:
    """Return True if the document already contains a Table of Contents."""
    for para in doc.paragraphs:
        style_name = (para.style.name if para.style else "").lower()
        # TOC-styled paragraphs
        if style_name.startswith("toc"):
            return True
        # Literal heading text
        text_lower = para.text.strip().lower()
        if text_lower in ("table of contents", "contents"):
            return True
    # Raw TOC field in XML
    body_xml = doc.element.body.xml
    if "TOC" in body_xml and "instrText" in body_xml:
        if re.search(r'<w:instrText[^>]*>\s*TOC\b', body_xml):
            return True
    return False


def _find_toc_insertion_point(doc: Document) -> int:
    """
    Return the index of the first real Heading 1 paragraph (after the title page).
    The TOC will be inserted immediately before this paragraph.
    Falls back to index 0 if no Heading 1 is found.
    """
    for i, para in enumerate(doc.paragraphs):
        style_name = (para.style.name if para.style else "").lower()
        if style_name == "heading 1" and para.text.strip():
            return i
    return 0


def _insert_toc_field(doc: Document, before_index: int) -> None:
    """
    Insert a TOC heading + real Word TOC field + page break immediately before
    the paragraph at *before_index*.  The TOC field auto-populates headings 1–3
    with hyperlinks and page numbers when the user opens the file in Word and
    accepts the 'Update Field' prompt (or presses F9 / Ctrl+A, F9). Combined
    with _force_update_fields_on_open, Word also rebuilds it automatically on open.
    """
    h_font = REQUIRED_HEADING_FONT or "Arial"
    h_size = HEADING_SIZE_BY_LEVEL.get(1, REQUIRED_HEADING_SIZE_PT)

    anchor = doc.paragraphs[before_index]

    # ── 1. "Table of Contents" heading paragraph ───────────────────────────
    toc_heading_p = OxmlElement("w:p")
    anchor._p.addprevious(toc_heading_p)
    toc_heading_para = Paragraph(toc_heading_p, anchor._parent)
    try:
        toc_heading_para.style = doc.styles["Heading 1"]
    except KeyError:
        pass
    run = toc_heading_para.add_run("Table of Contents")
    run.font.name = h_font
    run.font.size = Pt(h_size)
    run.bold = True

    # ── 2. Paragraph containing the TOC field ─────────────────────────────
    toc_para_p = OxmlElement("w:p")
    toc_heading_p.addnext(toc_para_p)

    r_begin = OxmlElement("w:r")
    fc_begin = OxmlElement("w:fldChar")
    fc_begin.set(qn("w:fldCharType"), "begin")
    fc_begin.set(qn("w:dirty"), "true")   # marks field as stale → Word updates on open
    r_begin.append(fc_begin)

    r_instr = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u \\n '
    r_instr.append(instr)

    r_sep = OxmlElement("w:r")
    fc_sep = OxmlElement("w:fldChar")
    fc_sep.set(qn("w:fldCharType"), "separate")
    r_sep.append(fc_sep)

    r_end = OxmlElement("w:r")
    fc_end = OxmlElement("w:fldChar")
    fc_end.set(qn("w:fldCharType"), "end")
    r_end.append(fc_end)

    toc_para_p.extend([r_begin, r_instr, r_sep, r_end])

    # ── 3. Page break paragraph after the TOC field ───────────────────────
    break_p = OxmlElement("w:p")
    toc_para_p.addnext(break_p)
    break_para = Paragraph(break_p, anchor._parent)
    break_para.add_run().add_break(WD_BREAK.PAGE)


def _force_update_fields_on_open(output_path: Path) -> None:
    """
    Inject <w:updateFields w:val="true"/> into word/settings.xml so Word
    automatically rebuilds all fields (including the TOC) the moment the
    file is opened — no manual right-click or F9 needed.
    """
    import io
    import zipfile as _zipfile

    with _zipfile.ZipFile(str(output_path), 'r') as zin:
        files = {name: zin.read(name) for name in zin.namelist()}

    if "word/settings.xml" not in files:
        return

    settings_xml = files["word/settings.xml"].decode("utf-8")
    if "w:updateFields" not in settings_xml:
        settings_xml = re.sub(
            r'(<w:settings\b[^>]*>)',
            r'\1<w:updateFields w:val="true"/>',
            settings_xml,
            count=1,
        )
        files["word/settings.xml"] = settings_xml.encode("utf-8")

        buf = io.BytesIO()
        with _zipfile.ZipFile(buf, 'w', _zipfile.ZIP_DEFLATED) as zout:
            for name, data in files.items():
                zout.writestr(name, data)
        output_path.write_bytes(buf.getvalue())


def add_missing_table_of_contents(doc: Document) -> None:
    """
    Public entry point.  No-ops if a TOC already exists; otherwise inserts
    a Heading 1 'Table of Contents' + a real Word TOC field (\\o 1-3 \\h \\z \\u)
    + a page break, placed immediately before the first Heading 1 content paragraph.
    """
    if _document_has_toc(doc):
        return
    insertion_idx = _find_toc_insertion_point(doc)
    if insertion_idx >= len(doc.paragraphs):
        return
    _insert_toc_field(doc, insertion_idx)


def build_final_fixed_doc(original_file: Path, state, output_path: Path):
    """
    Build the Final Fixed document.

    Order of operations:
    1. Copy original
    2. Strip ALL existing comments (pipeline must not add any — Final Fixed is clean)
    3. Clean spacing (in-place, preserving formatting)
    4. Remove close-range duplicate paragraphs
    5. Apply LLM rewrites (in-place)
    6. Add missing summary headings
    7. Add missing image captions and source links
    8. Insert generated quizzes at unit ends (drift-safe, one per unit, only if missing)
    9. Insert advanced visuals / diagrams (page-width capped, centred)
    10. Polish structure (remove excess blanks)
    11. Justify body paragraphs (never overriding intentional CENTER/RIGHT/DISTRIBUTE)
    12. Normalize styles and heading sizes (Heading 1 = 12pt, Heading 2 = 10pt, Heading 3 = 10pt, bold pseudo-headings = 14pt)
    13. Insert Table of Contents if missing
    14. Force Word to auto-update fields (TOC page numbers) on open
    """
    doc = copy_docx(original_file)
    if not getattr(state, "_cached_image_sources", None):
        _inject_image_issues(state, original_file)

    # ── Step 2: Strip ALL comments — Final Fixed must be completely clean ──
    _strip_all_comments(doc)

    # ── Step 3: Spacing ───────────────────────────────────────────────────
    clean_doc_spacing(doc)

    # ── Step 4: Remove close-range duplicate paragraphs ───────────────────
    remove_exact_duplicate_body_paragraphs(doc)

    # ── Step 5: LLM rewrites ──────────────────────────────────────────────
    apply_rewrites_to_doc(doc, state)

    # ── Step 6: Summary headings ──────────────────────────────────────────
    add_missing_summary_headings(doc)

    # ── Step 7: Image captions and source links ───────────────────────────
    # Reuse the same image_sources dict cached by _inject_image_issues so
    # the Fixed doc and Review Comments doc are guaranteed to be in sync.
    from image_source_finder import find_sources_for_images
    if not getattr(state, '_cached_image_sources', None):
        state._cached_image_sources = find_sources_for_images(state)
    image_sources = state._cached_image_sources

    figure_counter = _get_next_figure_number(doc)
    caption_inserted_at = set()
    source_inserted_at = set()  # tracks (target_idx, source_url) pairs already written

    for img in state.images:
        if img.paragraph_index is None:
            continue
        target_idx = (
            img.nearest_caption_paragraph_index
            if img.nearest_caption_paragraph_index is not None
            else img.paragraph_index
        )
        _img_source_url = image_sources.get(img.filename)
        needs_caption = (not img.nearest_caption_text) and (target_idx not in caption_inserted_at)
        needs_source  = (not img.has_source_link)    and ((target_idx, _img_source_url) not in source_inserted_at)

        if needs_caption or needs_source:
            figure_text = (
                f"Figure {figure_counter}"
                if needs_caption else None
            )
            source_text = _img_source_url if needs_source else None

            _append_caption_then_source(doc, target_idx, figure_text, source_text)

            if needs_caption:
                caption_inserted_at.add(target_idx)
                figure_counter += 1
            if needs_source:
                source_inserted_at.add((target_idx, _img_source_url))

    # ── Step 8: Quizzes (drift-safe, one per unit, only if missing) ───────
    insert_generated_quizzes_after_units(doc, state)

    # ── Step 9: Visuals / diagrams (page-width capped, centred) ──────────
    if getattr(state, "generated_visuals", None):
        insert_advanced_visuals(doc, state)
    elif getattr(state, "generated_diagrams", None):
        insert_generated_diagrams(doc, state)

    # ── Step 10: Polish ───────────────────────────────────────────────────
    polish_doc_structure(doc)

    # ── Step 11: Justify (preserves intentional CENTER/RIGHT/DISTRIBUTE) ──
    justify_body_paragraphs(doc)

    # ── Step 11.5: Final quiz alignment protection ────────────────────────
    force_quiz_alignment(doc)

    # ── Step 12: Normalize heading sizes (12/10/10) and document fonts ────
    normalize_heading_levels(doc)
    normalize_styles(doc)
    normalize_pseudo_heading_sizes(doc)

    # ── Step 13: Insert TOC if missing ───────────────────────────────────
    add_missing_table_of_contents(doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    # ── Step 14: Force Word to auto-update TOC on open ───────────────────
    _force_update_fields_on_open(output_path)
