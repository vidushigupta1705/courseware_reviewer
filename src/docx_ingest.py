import os
import re
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from models import ParagraphInfo, ImageInfo, DocumentState, TableInfo, TableCellInfo
from utils import pt_to_float, normalize_text, has_extra_spacing


ALIGNMENT_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: "LEFT",
    WD_ALIGN_PARAGRAPH.CENTER: "CENTER",
    WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY",
}


def detect_heading_level(style_name: str) -> Optional[int]:
    if not style_name:
        return None
    style_name = style_name.strip().lower()
    if style_name.startswith("heading"):
        parts = style_name.split()
        if len(parts) >= 2 and parts[1].isdigit():
            level = int(parts[1])
            if 1 <= level <= 9:
                return level
    return None


def is_code_like_paragraph(text: str, style_name: str = "") -> bool:
    if not text or not text.strip():
        return False

    lower_style = (style_name or "").lower()
    if any(x in lower_style for x in ["code", "source code", "preformatted"]):
        return True

    score = 0
    if re.search(r"[{}();=<>\[\]#]", text):
        score += 1
    if re.search(r"\bdef\s+\w+\s*\(", text):
        score += 2
    if re.search(r"\bclass\s+\w+", text):
        score += 2
    if re.search(r"\bimport\s+\w+", text):
        score += 2
    if re.search(r"\bfor\b.*:", text):
        score += 1
    if re.search(r"\bif\b.*:", text):
        score += 1
    if re.search(r"^\s{2,}\S+", text, flags=re.MULTILINE):
        score += 1
    if re.search(r"<[^>]+>", text):
        score += 1

    return score >= 3


# ── Heading heuristic helpers ─────────────────────────────────────────────────

def _collect_body_font_sizes(doc: Document) -> List[float]:
    sizes = []
    
    # First try run-level explicit sizes
    for para in doc.paragraphs:
        style_name = (para.style.name if para.style else "Normal").lower()
        if style_name.startswith("heading"):
            continue
        for run in para.runs:
            if run.text.strip() and run.font.size:
                val = pt_to_float(run.font.size)
                if val is not None:
                    sizes.append(val)

    # Fallback: read from the Normal style definition if runs had no explicit size
    if not sizes:
        try:
            normal_style = doc.styles["Normal"]
            if normal_style.font and normal_style.font.size:
                val = pt_to_float(normal_style.font.size)
                if val is not None:
                    sizes = [val]
        except Exception:
            pass

    # Final fallback: standard body size
    if not sizes:
        sizes = [10.0]

    return sizes


def _infer_heading_level_heuristic(para, body_font_sizes: List[float]) -> Optional[int]:
    """
    Fallback heading detection for documents that use bold + large text
    instead of proper Word heading styles.

    All three conditions must be true:
      - Text is short (<=12 words)
      - Every non-empty run is explicitly bold
      - Average font size is larger than the typical body size

    Returns a synthetic heading level (1-3) or None.
    """
    text = para.text.strip()
    if not text or len(text.split()) > 12:
        return None

    runs_with_text = [r for r in para.runs if r.text.strip()]
    if not runs_with_text:
        return None

    # Every run must be explicitly bold
    def _is_bold(run) -> bool:
        if run.bold is True:
            return True
        # Check if bold is inherited from the paragraph style
        if run.bold is None and run.style and run.style.font:
            return bool(run.style.font.bold)
        return False

    if not all(_is_bold(r) for r in runs_with_text):
        return None

    sizes = [r.font.size.pt for r in runs_with_text if r.font.size]
    if not sizes:
        return None

    avg_size = sum(sizes) / len(sizes)
    typical_body = (sum(body_font_sizes) / len(body_font_sizes)) if body_font_sizes else 10.0

    if avg_size >= typical_body + 4:
        return 1
    if avg_size >= typical_body + 2:
        return 2
    if avg_size > typical_body:
        return 3
    return None


# ── Main extraction functions ─────────────────────────────────────────────────

def extract_paragraphs(doc: Document) -> List[ParagraphInfo]:
    paragraphs = []
    body_font_sizes = _collect_body_font_sizes(doc)

    for idx, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else "Normal"
        heading_level = detect_heading_level(style_name)
        is_heading = heading_level is not None

        # Fallback: catch headings styled with bold + large font but no Word heading style
        if not is_heading:
            heuristic_level = _infer_heading_level_heuristic(para, body_font_sizes)
            if heuristic_level is not None:
                heading_level = heuristic_level
                is_heading = True

        font_names = []
        font_sizes = []
        bold_flags = []

        for run in para.runs:
            if run.text.strip():
                if run.font.name:
                    font_names.append(run.font.name)
                elif run.style and run.style.font and run.style.font.name:
                    font_names.append(run.style.font.name)

                if run.font.size:
                    size_val = pt_to_float(run.font.size)
                    if size_val is not None:
                        font_sizes.append(size_val)

                if run.bold is not None:
                    bold_flags.append(bool(run.bold))

        text = normalize_text(para.text)

        paragraph_info = ParagraphInfo(
            index=idx,
            text=text,
            style_name=style_name,
            is_heading=is_heading,
            heading_level=heading_level,
            font_names=font_names,
            font_sizes=font_sizes,
            bold_flags=bold_flags,
            alignment=ALIGNMENT_MAP.get(para.alignment, "UNKNOWN"),
            has_extra_spacing_issue=has_extra_spacing(para.text),
            is_code_like=is_code_like_paragraph(para.text, style_name),
        )
        paragraphs.append(paragraph_info)

    return paragraphs


def extract_tables(doc: Document) -> List[TableInfo]:
    tables = []

    W_P = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"
    W_TBL = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl"

    para_count = 0
    table_para_positions = []

    for child in doc.element.body:
        if child.tag == W_P:
            para_count += 1
        elif child.tag == W_TBL:
            table_para_positions.append(max(0, para_count - 1))

    for t_idx, table in enumerate(doc.tables):
        row_count = len(table.rows)
        col_count = max((len(r.cells) for r in table.rows), default=0)

        cell_items = []
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell_items.append(
                    TableCellInfo(
                        row_index=r_idx,
                        col_index=c_idx,
                        text=normalize_text(cell.text),
                    )
                )

        para_idx = table_para_positions[t_idx] if t_idx < len(table_para_positions) else None

        tables.append(
            TableInfo(
                table_index=t_idx,
                row_count=row_count,
                col_count=col_count,
                cells=cell_items,
                paragraph_index=para_idx,
            )
        )

    return tables


def _get_rel_ids_from_paragraph(paragraph):
    rel_ids = []
    drawings = paragraph._element.xpath('.//*[local-name()="drawing"]')
    for drawing in drawings:
        blips = drawing.xpath('.//*[local-name()="blip"]')
        for blip in blips:
            for attr_name, attr_value in blip.attrib.items():
                if attr_name.endswith('}embed'):
                    rel_ids.append(attr_value)
    return rel_ids


def build_image_paragraph_map(doc: Document):
    """
    Returns dict: rel_id -> paragraph index
    Scans both regular paragraphs and table cells for embedded images.
    """
    rel_to_para = {}

    for p_idx, para in enumerate(doc.paragraphs):
        rel_ids = _get_rel_ids_from_paragraph(para)
        if rel_ids:
            for rel_id in rel_ids:
                if rel_id not in rel_to_para:
                    rel_to_para[rel_id] = p_idx

    W_P = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"
    W_TBL = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl"

    para_count = 0
    for child in doc.element.body:
        if child.tag == W_P:
            para_count += 1
        elif child.tag == W_TBL:
            anchor_idx = max(0, para_count - 1)
            drawings = child.xpath('.//*[local-name()="drawing"]')
            for drawing in drawings:
                blips = drawing.xpath('.//*[local-name()="blip"]')
                for blip in blips:
                    for attr_name, attr_value in blip.attrib.items():
                        if attr_name.endswith('}embed'):
                            rel_id = attr_value
                            if rel_id not in rel_to_para:
                                rel_to_para[rel_id] = anchor_idx

    return rel_to_para


def extract_images(doc: Document, output_dir: Path) -> List[ImageInfo]:
    output_dir.mkdir(parents=True, exist_ok=True)
    images = []
    rels = doc.part.rels
    image_para_map = build_image_paragraph_map(doc)

    seen = set()

    for rel_id, rel in rels.items():
        if "image" not in rel.target_ref:
            continue

        if rel_id not in image_para_map:
            continue

        image_part = rel.target_part
        image_bytes = image_part.blob

        if len(image_bytes) < 1024:
            continue

        original_name = os.path.basename(rel.target_ref)
        if not original_name:
            original_name = f"{rel_id}.png"

        output_path = output_dir / original_name
        candidate = output_path
        count = 1
        while candidate.exists() or str(candidate) in seen:
            stem = output_path.stem
            suffix = output_path.suffix
            candidate = output_dir / f"{stem}_{count}{suffix}"
            count += 1

        with open(candidate, "wb") as f:
            f.write(image_bytes)

        seen.add(str(candidate))
        images.append(
            ImageInfo(
                rel_id=rel_id,
                filename=os.path.basename(candidate),
                output_path=str(candidate),
                paragraph_index=image_para_map.get(rel_id),
            )
        )

    return images


def load_docx_state(input_file: Path, extracted_images_dir: Path) -> DocumentState:
    doc = Document(str(input_file))
    base_filename = input_file.stem

    paragraphs = extract_paragraphs(doc)
    tables = extract_tables(doc)
    images = extract_images(doc, extracted_images_dir / base_filename)

    return DocumentState(
        input_file=str(input_file),
        base_filename=base_filename,
        paragraphs=paragraphs,
        tables=tables,
        images=images,
        ocr_results=[],
        rewrite_suggestions=[],
        units=[],
        issues=[],
    )
